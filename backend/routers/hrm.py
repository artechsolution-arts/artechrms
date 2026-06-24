import time
from calendar import monthrange as _monthrange
from datetime import date, timedelta
from typing import Optional
from backend import storage

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query, Request
from backend.approval_utils import require_approval_rights
from pydantic import BaseModel
from sqlalchemy.orm import Session

from backend.database import get_db
from backend.models.auth import User
from backend.models.employee import Employee
from backend.models.leave import LeaveType
from backend.models.hrm import (
    Announcement,
    EmergencyContact,
    EmployeeAsset,
    EmployeeDocument,
    EmployeeHistory,
    ExpenseClaim,
    Holiday,
    LeaveBalance,
)

router = APIRouter(prefix="/api/hrm", tags=["HRM"])

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _employee_or_404(emp_id: int, db: Session) -> Employee:
    emp = db.query(Employee).filter(Employee.id == emp_id).first()
    if not emp:
        raise HTTPException(status_code=404, detail="Employee not found")
    return emp


# ===========================================================================
# 1. EMERGENCY CONTACTS
# ===========================================================================

class EmergencyContactIn(BaseModel):
    name: str
    relationship_type: str
    phone: str
    email: Optional[str] = None
    is_primary: bool = False


@router.get("/employees/{emp_id}/emergency-contacts")
def list_emergency_contacts(emp_id: int, db: Session = Depends(get_db)):
    _employee_or_404(emp_id, db)
    contacts = (
        db.query(EmergencyContact)
        .filter(EmergencyContact.employee_id == emp_id)
        .all()
    )
    return [
        {
            "id": c.id,
            "employee_id": c.employee_id,
            "name": c.name,
            "relationship_type": c.relationship_type,
            "phone": c.phone,
            "email": c.email,
            "is_primary": c.is_primary,
            "created_at": str(c.created_at) if c.created_at else None,
        }
        for c in contacts
    ]


@router.post("/employees/{emp_id}/emergency-contacts", status_code=201)
def create_emergency_contact(
    emp_id: int, data: EmergencyContactIn, db: Session = Depends(get_db)
):
    _employee_or_404(emp_id, db)
    contact = EmergencyContact(employee_id=emp_id, **data.model_dump())
    db.add(contact)
    db.commit()
    db.refresh(contact)
    return {"id": contact.id, "ok": True}


@router.put("/emergency-contacts/{contact_id}")
def update_emergency_contact(
    contact_id: int, data: EmergencyContactIn, db: Session = Depends(get_db)
):
    contact = db.query(EmergencyContact).filter(EmergencyContact.id == contact_id).first()
    if not contact:
        raise HTTPException(status_code=404, detail="Emergency contact not found")
    for field, value in data.model_dump(exclude_unset=True).items():
        setattr(contact, field, value)
    db.commit()
    return {"ok": True}


@router.delete("/emergency-contacts/{contact_id}")
def delete_emergency_contact(contact_id: int, db: Session = Depends(get_db)):
    contact = db.query(EmergencyContact).filter(EmergencyContact.id == contact_id).first()
    if not contact:
        raise HTTPException(status_code=404, detail="Emergency contact not found")
    db.delete(contact)
    db.commit()
    return {"ok": True}


# ===========================================================================
# 2. EMPLOYEE DOCUMENTS
# ===========================================================================

ALLOWED_DOC_TYPES = {
    "Aadhaar", "PAN", "Passport", "Offer Letter",
    "Experience Letter", "Education Certificate", "Other",
}


@router.get("/employees/{emp_id}/documents")
def list_documents(emp_id: int, db: Session = Depends(get_db)):
    _employee_or_404(emp_id, db)
    docs = (
        db.query(EmployeeDocument)
        .filter(EmployeeDocument.employee_id == emp_id)
        .order_by(EmployeeDocument.uploaded_at.desc())
        .all()
    )
    return [
        {
            "id": d.id,
            "employee_id": d.employee_id,
            "document_type": d.document_type,
            "document_name": d.document_name,
            "file_url": d.file_url,
            "uploaded_at": str(d.uploaded_at) if d.uploaded_at else None,
        }
        for d in docs
    ]


@router.post("/employees/{emp_id}/documents", status_code=201)
async def upload_document(
    emp_id: int,
    document_type: str = Query(..., description="Document type"),
    document_name: str = Query(..., description="Human-readable document name"),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    _employee_or_404(emp_id, db)

    if document_type not in ALLOWED_DOC_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"document_type must be one of: {', '.join(sorted(ALLOWED_DOC_TYPES))}",
        )

    ext = (file.filename or "file").rsplit(".", 1)[-1].lower()
    fname = f"doc_{emp_id}_{int(time.time())}.{ext}"
    file_url = storage.upload_file(await file.read(), "documents", fname)
    doc = EmployeeDocument(
        employee_id=emp_id,
        document_type=document_type,
        document_name=document_name,
        file_url=file_url,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return {"id": doc.id, "file_url": file_url, "ok": True}


@router.delete("/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(EmployeeDocument).filter(EmployeeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    storage.delete_file(doc.file_url)
    db.delete(doc)
    db.commit()
    return {"ok": True}


# ===========================================================================
# 3. HOLIDAYS
# ===========================================================================

class HolidayIn(BaseModel):
    name: str
    date: date
    holiday_type: str = "Public"
    description: Optional[str] = None


@router.get("/holidays")
def list_holidays(year: Optional[int] = None, db: Session = Depends(get_db)):
    q = db.query(Holiday)
    if year:
        from sqlalchemy import extract
        q = q.filter(extract("year", Holiday.date) == year)
    holidays = q.order_by(Holiday.date).all()
    return [
        {
            "id": h.id,
            "name": h.name,
            "date": str(h.date),
            "holiday_type": h.holiday_type,
            "description": h.description,
            "created_at": str(h.created_at) if h.created_at else None,
        }
        for h in holidays
    ]


@router.post("/holidays", status_code=201)
def create_holiday(data: HolidayIn, db: Session = Depends(get_db)):
    holiday = Holiday(**data.model_dump())
    db.add(holiday)
    db.commit()
    db.refresh(holiday)
    return {"id": holiday.id, "ok": True}


@router.delete("/holidays/{holiday_id}")
def delete_holiday(holiday_id: int, db: Session = Depends(get_db)):
    holiday = db.query(Holiday).filter(Holiday.id == holiday_id).first()
    if not holiday:
        raise HTTPException(status_code=404, detail="Holiday not found")
    db.delete(holiday)
    db.commit()
    return {"ok": True}


# ===========================================================================
# 4. LEAVE BALANCES
# ===========================================================================

class LeaveBalanceAllocateIn(BaseModel):
    employee_id: int
    leave_type_id: int
    year: int
    allocated: float


class AllocateAllIn(BaseModel):
    year: int


@router.get("/leave-balances")
def list_leave_balances(
    employee_id: Optional[int] = None,
    year: Optional[int] = None,
    db: Session = Depends(get_db),
):
    q = db.query(LeaveBalance)
    if employee_id:
        q = q.filter(LeaveBalance.employee_id == employee_id)
    if year:
        q = q.filter(LeaveBalance.year == year)
    balances = q.all()
    result = []
    for b in balances:
        result.append({
            "id": b.id,
            "employee_id": b.employee_id,
            "employee_name": b.employee_rel.full_name if b.employee_rel else None,
            "profile_photo": b.employee_rel.profile_photo if b.employee_rel else None,
            "leave_type_id": b.leave_type_id,
            "leave_type": b.leave_type_rel.name if b.leave_type_rel else None,
            "year": b.year,
            "allocated": b.allocated,
            "used": b.used,
            "carried_forward": b.carried_forward,
            "available": round(b.allocated + b.carried_forward - b.used, 2),
        })
    return result


@router.post("/leave-balances/allocate", status_code=201)
def allocate_leave_balance(data: LeaveBalanceAllocateIn, db: Session = Depends(get_db)):
    balance = (
        db.query(LeaveBalance)
        .filter(
            LeaveBalance.employee_id == data.employee_id,
            LeaveBalance.leave_type_id == data.leave_type_id,
            LeaveBalance.year == data.year,
        )
        .first()
    )
    if balance:
        balance.allocated = data.allocated
    else:
        balance = LeaveBalance(**data.model_dump())
        db.add(balance)
    db.commit()
    db.refresh(balance)
    return {"id": balance.id, "ok": True}


@router.post("/leave-balances/allocate-all", status_code=201)
def allocate_all_leave_balances(data: AllocateAllIn, db: Session = Depends(get_db)):
    """Initialize leave balance records for active employees.
    Only creates missing records (allocated=0); never overwrites HR-set balances.
    HR should manually set each employee's allocated balance via the edit modal.
    """
    employees = db.query(Employee).filter(Employee.status == "Active").all()
    leave_types = db.query(LeaveType).all()
    created = 0
    skipped = 0
    for emp in employees:
        for lt in leave_types:
            balance = (
                db.query(LeaveBalance)
                .filter(
                    LeaveBalance.employee_id == emp.id,
                    LeaveBalance.leave_type_id == lt.id,
                    LeaveBalance.year == data.year,
                )
                .first()
            )
            if balance:
                skipped += 1  # already set by HR — do not overwrite
            else:
                balance = LeaveBalance(
                    employee_id=emp.id,
                    leave_type_id=lt.id,
                    year=data.year,
                    allocated=0,
                    used=0,
                    carried_forward=0,
                )
                db.add(balance)
                created += 1
    db.commit()
    return {"ok": True, "created": created, "skipped": skipped, "year": data.year}


@router.get("/leave-balances/employee/{emp_id}")
def get_employee_leave_balances(emp_id: int, db: Session = Depends(get_db)):
    _employee_or_404(emp_id, db)
    balances = db.query(LeaveBalance).filter(LeaveBalance.employee_id == emp_id).all()
    return [
        {
            "id": b.id,
            "leave_type_id": b.leave_type_id,
            "leave_type": b.leave_type_rel.name if b.leave_type_rel else None,
            "year": b.year,
            "allocated": b.allocated,
            "used": b.used,
            "carried_forward": b.carried_forward,
            "available": round(b.allocated + b.carried_forward - b.used, 2),
        }
        for b in balances
    ]


# ===========================================================================
# 5. ANNOUNCEMENTS
# ===========================================================================

class AnnouncementIn(BaseModel):
    title: str
    content: str
    priority: str = "Medium"
    expires_on: Optional[date] = None
    created_by: Optional[str] = None


class AnnouncementUpdateIn(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    priority: Optional[str] = None
    is_active: Optional[bool] = None
    expires_on: Optional[date] = None


@router.get("/announcements")
def list_announcements(
    active_only: bool = False,
    db: Session = Depends(get_db),
):
    q = db.query(Announcement)
    if active_only:
        q = q.filter(Announcement.is_active == True)  # noqa: E712
    announcements = q.order_by(Announcement.created_at.desc()).all()
    return [
        {
            "id": a.id,
            "title": a.title,
            "content": a.content,
            "priority": a.priority,
            "is_active": a.is_active,
            "created_by": a.created_by,
            "created_at": str(a.created_at) if a.created_at else None,
            "expires_on": str(a.expires_on) if a.expires_on else None,
        }
        for a in announcements
    ]


@router.post("/announcements", status_code=201)
def create_announcement(data: AnnouncementIn, db: Session = Depends(get_db)):
    ann = Announcement(**data.model_dump())
    db.add(ann)
    db.commit()
    db.refresh(ann)
    return {"id": ann.id, "ok": True}


@router.put("/announcements/{ann_id}")
def update_announcement(
    ann_id: int, data: AnnouncementUpdateIn, db: Session = Depends(get_db)
):
    ann = db.query(Announcement).filter(Announcement.id == ann_id).first()
    if not ann:
        raise HTTPException(status_code=404, detail="Announcement not found")
    for field, value in data.model_dump(exclude_unset=True).items():
        setattr(ann, field, value)
    db.commit()
    return {"ok": True}


@router.delete("/announcements/{ann_id}")
def delete_announcement(ann_id: int, db: Session = Depends(get_db)):
    ann = db.query(Announcement).filter(Announcement.id == ann_id).first()
    if not ann:
        raise HTTPException(status_code=404, detail="Announcement not found")
    db.delete(ann)
    db.commit()
    return {"ok": True}


# ===========================================================================
# 6. EXPENSE CLAIMS
# ===========================================================================

class ExpenseClaimIn(BaseModel):
    employee_id: int
    claim_date: date
    expense_type: str
    amount: float
    description: Optional[str] = None


class ExpenseActionIn(BaseModel):
    remarks: Optional[str] = None


@router.get("/expenses")
def list_expenses(
    employee_id: Optional[int] = None,
    status: Optional[str] = None,
    db: Session = Depends(get_db),
):
    q = db.query(ExpenseClaim)
    if employee_id:
        q = q.filter(ExpenseClaim.employee_id == employee_id)
    if status:
        q = q.filter(ExpenseClaim.status == status)
    expenses = q.order_by(ExpenseClaim.claim_date.desc()).all()
    return [
        {
            "id": e.id,
            "employee_id": e.employee_id,
            "employee_name": e.employee_rel.full_name if e.employee_rel else None,
            "claim_date": str(e.claim_date),
            "expense_type": e.expense_type,
            "amount": e.amount,
            "description": e.description,
            "receipt_url": e.receipt_url,
            "status": e.status,
            "approved_by": e.approved_by,
            "approved_on": str(e.approved_on) if e.approved_on else None,
            "remarks": e.remarks,
            "created_at": str(e.created_at) if e.created_at else None,
        }
        for e in expenses
    ]


@router.post("/expenses", status_code=201)
def create_expense(data: ExpenseClaimIn, db: Session = Depends(get_db)):
    _employee_or_404(data.employee_id, db)
    expense = ExpenseClaim(**data.model_dump())
    db.add(expense)
    db.commit()
    db.refresh(expense)
    return {"id": expense.id, "ok": True}


@router.put("/expenses/{expense_id}/approve")
def approve_expense(
    expense_id: int, data: ExpenseActionIn, db: Session = Depends(get_db)
):
    expense = db.query(ExpenseClaim).filter(ExpenseClaim.id == expense_id).first()
    if not expense:
        raise HTTPException(status_code=404, detail="Expense claim not found")
    expense.status = "Approved"
    expense.remarks = data.remarks
    expense.approved_on = date.today()
    db.commit()
    _notify_expense_result(db, expense, "Approved")
    return {"ok": True}


@router.put("/expenses/{expense_id}/reject")
def reject_expense(
    expense_id: int, data: ExpenseActionIn, db: Session = Depends(get_db)
):
    expense = db.query(ExpenseClaim).filter(ExpenseClaim.id == expense_id).first()
    if not expense:
        raise HTTPException(status_code=404, detail="Expense claim not found")
    expense.status = "Rejected"
    expense.remarks = data.remarks
    expense.approved_on = date.today()
    db.commit()
    _notify_expense_result(db, expense, "Rejected")
    return {"ok": True}


def _notify_expense_result(db, expense, status: str) -> None:
    """In-app notification to the employee when their expense claim is actioned."""
    try:
        emp = db.query(Employee).filter(Employee.id == expense.employee_id).first()
        if not emp or not emp.user_id:
            return
        from backend.services.notification_service import push
        push(
            db, emp.user_id, "expense",
            f"Expense Claim {status}",
            f"Your ₹{int(expense.amount):,} {expense.expense_type} claim has been {status.lower()}"
            + (f": {expense.remarks}" if expense.remarks else "."),
            entity_id=expense.id,
            notif_type="approval_result",
            action="emp-expenses",
            priority="high" if status == "Rejected" else "medium",
        )
        db.commit()
    except Exception:
        pass


@router.delete("/expenses/{expense_id}")
def delete_expense(expense_id: int, db: Session = Depends(get_db)):
    expense = db.query(ExpenseClaim).filter(ExpenseClaim.id == expense_id).first()
    if not expense:
        raise HTTPException(status_code=404, detail="Expense claim not found")
    db.delete(expense)
    db.commit()
    return {"ok": True}


# ===========================================================================
# 7. EMPLOYEE ASSETS
# ===========================================================================

class AssetIn(BaseModel):
    employee_id: int
    asset_name: str
    asset_type: str
    serial_number: Optional[str] = None
    allocated_date: date
    condition: str = "Good"
    notes: Optional[str] = None


class AssetReturnIn(BaseModel):
    returned_date: date
    condition: str = "Good"
    notes: Optional[str] = None


@router.get("/assets")
def list_assets(
    employee_id: Optional[int] = None,
    status: Optional[str] = None,
    db: Session = Depends(get_db),
):
    q = db.query(EmployeeAsset)
    if employee_id:
        q = q.filter(EmployeeAsset.employee_id == employee_id)
    if status:
        q = q.filter(EmployeeAsset.status == status)
    assets = q.order_by(EmployeeAsset.allocated_date.desc()).all()
    return [
        {
            "id": a.id,
            "employee_id": a.employee_id,
            "employee_name": a.employee_rel.full_name if a.employee_rel else None,
            "asset_name": a.asset_name,
            "asset_type": a.asset_type,
            "serial_number": a.serial_number,
            "allocated_date": str(a.allocated_date),
            "returned_date": str(a.returned_date) if a.returned_date else None,
            "condition": a.condition,
            "notes": a.notes,
            "status": a.status,
            "created_at": str(a.created_at) if a.created_at else None,
        }
        for a in assets
    ]


@router.post("/assets", status_code=201)
def create_asset(data: AssetIn, db: Session = Depends(get_db)):
    _employee_or_404(data.employee_id, db)
    asset = EmployeeAsset(**data.model_dump())
    db.add(asset)
    db.commit()
    db.refresh(asset)
    return {"id": asset.id, "ok": True}


@router.put("/assets/{asset_id}/return")
def return_asset(
    asset_id: int, data: AssetReturnIn, db: Session = Depends(get_db)
):
    asset = db.query(EmployeeAsset).filter(EmployeeAsset.id == asset_id).first()
    if not asset:
        raise HTTPException(status_code=404, detail="Asset not found")
    asset.returned_date = data.returned_date
    asset.condition = data.condition
    asset.notes = data.notes
    asset.status = "Returned"
    db.commit()
    return {"ok": True}


@router.delete("/assets/{asset_id}")
def delete_asset(asset_id: int, db: Session = Depends(get_db)):
    asset = db.query(EmployeeAsset).filter(EmployeeAsset.id == asset_id).first()
    if not asset:
        raise HTTPException(status_code=404, detail="Asset not found")
    db.delete(asset)
    db.commit()
    return {"ok": True}


# ===========================================================================
# 8. EMPLOYEE HISTORY
# ===========================================================================

class EmployeeHistoryIn(BaseModel):
    change_type: str = "Transfer"
    from_department: Optional[str] = None
    to_department: Optional[str] = None
    from_designation: Optional[str] = None
    to_designation: Optional[str] = None
    effective_date: date
    salary_before: Optional[float] = None
    salary_after: Optional[float] = None
    last_working_date: Optional[date] = None
    remarks: Optional[str] = None
    created_by: Optional[str] = None


@router.get("/employees/{emp_id}/history")
def get_employee_history(emp_id: int, db: Session = Depends(get_db)):
    _employee_or_404(emp_id, db)
    records = (
        db.query(EmployeeHistory)
        .filter(EmployeeHistory.employee_id == emp_id)
        .order_by(EmployeeHistory.effective_date.desc())
        .all()
    )
    return [
        {
            "id": r.id,
            "employee_id": r.employee_id,
            "change_type": r.change_type,
            "from_department": r.from_department,
            "to_department": r.to_department,
            "from_designation": r.from_designation,
            "to_designation": r.to_designation,
            "effective_date": str(r.effective_date),
            "salary_before": r.salary_before,
            "salary_after": r.salary_after,
            "last_working_date": str(r.last_working_date) if r.last_working_date else None,
            "remarks": r.remarks,
            "created_by": r.created_by,
            "created_at": str(r.created_at) if r.created_at else None,
            "updated_at": str(r.updated_at) if getattr(r, 'updated_at', None) else None,
        }
        for r in records
    ]


@router.post("/employees/{emp_id}/history", status_code=201)
def add_employee_history(
    emp_id: int, data: EmployeeHistoryIn, db: Session = Depends(get_db)
):
    _employee_or_404(emp_id, db)
    record = EmployeeHistory(employee_id=emp_id, **data.model_dump())
    db.add(record)
    db.commit()
    db.refresh(record)
    return {"id": record.id, "ok": True}


@router.put("/employees/{emp_id}/history/{record_id}")
def update_employee_history(
    emp_id: int, record_id: int, data: EmployeeHistoryIn, db: Session = Depends(get_db)
):
    record = db.query(EmployeeHistory).filter(
        EmployeeHistory.id == record_id,
        EmployeeHistory.employee_id == emp_id,
    ).first()
    if not record:
        raise HTTPException(404, "History record not found")
    for field, val in data.model_dump(exclude_unset=True).items():
        setattr(record, field, val)
    db.commit()
    db.refresh(record)
    return {"id": record.id, "ok": True}


@router.delete("/employees/{emp_id}/history/{record_id}")
def delete_employee_history(emp_id: int, record_id: int, db: Session = Depends(get_db)):
    record = db.query(EmployeeHistory).filter(
        EmployeeHistory.id == record_id,
        EmployeeHistory.employee_id == emp_id,
    ).first()
    if not record:
        raise HTTPException(404, "History record not found")
    db.delete(record)
    db.commit()
    return {"ok": True}


# ===========================================================================
# DOCUMENT REQUESTS (HR side)
# ===========================================================================

from backend.models.document_request import DocumentRequest
from backend.models.employee import Employee as _Emp
from datetime import datetime as _dt


@router.get("/document-requests")
def list_document_requests(status: Optional[str] = None, db: Session = Depends(get_db)):
    q = db.query(DocumentRequest).join(_Emp, DocumentRequest.employee_id == _Emp.id)
    if status and status != "All":
        q = q.filter(DocumentRequest.status == status)
    rows = q.order_by(DocumentRequest.requested_at.desc()).all()
    return [
        {
            "id": r.id,
            "employee_id": r.employee_id,
            "employee_name": r.employee.full_name if r.employee else "—",
            "employee_code": r.employee.employee_id if r.employee else "—",
            "profile_photo": r.employee.profile_photo if r.employee else None,
            "doc_type": r.doc_type,
            "remarks": r.remarks,
            "status": r.status,
            "requested_at": str(r.requested_at)[:10],
            "fulfilled_at": str(r.fulfilled_at)[:10] if r.fulfilled_at else None,
            "file_url": r.file_url,
            "file_name": r.file_name,
        }
        for r in rows
    ]


@router.post("/document-requests/{req_id}/upload")
async def upload_document_request(req_id: int, file: UploadFile = File(...), db: Session = Depends(get_db)):
    req = db.query(DocumentRequest).filter(DocumentRequest.id == req_id).first()
    if not req:
        raise HTTPException(404, "Request not found")
    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in ("pdf", "doc", "docx", "jpg", "jpeg", "png"):
        raise HTTPException(400, "Allowed formats: PDF, DOC, DOCX, JPG, PNG")
    fname = f"req_{req_id}_{int(time.time())}.{ext}"
    req.file_url = storage.upload_file(await file.read(), "documents", fname)
    req.file_name = file.filename
    req.status = "Fulfilled"
    req.fulfilled_at = _dt.utcnow()
    db.commit()
    return {"file_url": req.file_url, "ok": True}


# ── Status Sheets (HR read-only view) ─────────────────────────

from backend.models.status_entry import StatusEntry as _StatusEntry


@router.get("/status")
def hr_get_status(employee_id: int, month: str, db: Session = Depends(get_db)):
    try:
        year, mon = map(int, month.split("-"))
    except Exception:
        raise HTTPException(400, "Use YYYY-MM format")

    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        raise HTTPException(404, "Employee not found")

    start = date(year, mon, 1)
    end = date(year, mon, _monthrange(year, mon)[1])

    entries = (
        db.query(_StatusEntry)
        .filter(
            _StatusEntry.employee_id == employee_id,
            _StatusEntry.entry_date >= start,
            _StatusEntry.entry_date <= end,
        )
        .order_by(_StatusEntry.entry_date.asc())
        .all()
    )

    return {
        "employee_name": emp.full_name,
        "employee_code": emp.employee_id,
        "entries": [
            {
                "id": e.id,
                "task_id": e.task_id,
                "entry_date": str(e.entry_date),
                "task_name": e.task_name,
                "due_date": str(e.due_date) if e.due_date else None,
                "status": e.status,
                "percent_complete": e.percent_complete or 0,
            }
            for e in entries
        ],
    }


from pydantic import BaseModel as _BM
from typing import Optional as _Opt
from datetime import date as _date_t


class _StatusUpdate(_BM):
    task_name:        _Opt[str]   = None
    due_date:         _Opt[str]   = None
    status:           _Opt[str]   = None
    percent_complete: _Opt[int]   = None


@router.put("/status/{entry_id}")
def hr_update_status_entry(entry_id: int, data: _StatusUpdate, db: Session = Depends(get_db)):
    entry = db.query(_StatusEntry).filter(_StatusEntry.id == entry_id).first()
    if not entry:
        raise HTTPException(404, "Entry not found")
    if data.task_name is not None:
        entry.task_name = data.task_name
    if data.due_date is not None:
        try:
            entry.due_date = _date_t.fromisoformat(data.due_date) if data.due_date else None
        except ValueError:
            pass
    if data.status is not None:
        entry.status = data.status
    if data.percent_complete is not None:
        entry.percent_complete = max(0, min(100, data.percent_complete))
    db.commit()
    return {"ok": True}


# ── Team Leaves calendar (HR / CEO view) ──────────────────────

@router.get("/team-leaves")
def hr_team_leaves(month: str, db: Session = Depends(get_db)):
    """All approved/pending leaves for every employee — used by HR & CEO work-mode calendar."""
    from backend.models.leave import LeaveApplication, LeaveType
    try:
        year, mon = map(int, month.split("-"))
    except Exception:
        raise HTTPException(400, "Use YYYY-MM format")
    from calendar import monthrange as _mr
    start = date(year, mon, 1)
    end   = date(year, mon, _mr(year, mon)[1])
    leaves = (
        db.query(LeaveApplication)
        .join(Employee, LeaveApplication.employee_id == Employee.id)
        .filter(
            LeaveApplication.status.in_(["Pending", "Approved"]),
            LeaveApplication.from_date <= end,
            LeaveApplication.to_date   >= start,
        )
        .order_by(LeaveApplication.from_date.asc())
        .all()
    )
    result = []
    for lv in leaves:
        emp = lv.employee_rel
        lt  = lv.leave_type_rel
        result.append({
            "id":             lv.id,
            "employee_name":  emp.full_name if emp else "",
            "profile_photo":  emp.profile_photo if emp else None,
            "leave_type":     lt.name if lt else "",
            "from_date":      str(lv.from_date),
            "to_date":        str(lv.to_date),
            "total_days":     lv.total_days,
            "half_day":       lv.half_day,
            "leave_category": lv.leave_category or "Planned",
            "status":         lv.status,
            "reason":         lv.reason,
        })
    return result


# ── Work Mode Sheet (HR view + approve/reject) ─────────────────

from datetime import datetime as _wm_dt
from backend.models.work_mode_entry import WorkModeEntry as _WMEntry
from pydantic import BaseModel as _BaseModel
from typing import Optional as _Optional


class _WMAction(_BaseModel):
    remarks: _Optional[str] = None


def _wm_hr_dict(e) -> dict:
    return {
        "id":            e.id,
        "employee_id":   e.employee.id if e.employee else None,
        "employee_name": e.employee.full_name if e.employee else "—",
        "employee_code": e.employee.employee_id if e.employee else "—",
        "profile_photo": e.employee.profile_photo if e.employee else None,
        "entry_date":    str(e.entry_date),
        "work_mode":     e.work_mode,
        "reason":        e.reason,
        "duration":      e.duration,
        "status":        e.status,
        "hr_remarks":    e.hr_remarks,
        "created_at":    str(e.created_at)[:10],
    }


@router.get("/work-mode")
def hr_list_work_mode(
    month: str,
    employee_id: _Optional[int] = None,
    status: _Optional[str] = None,
    db: Session = Depends(get_db),
):
    try:
        year, mon = map(int, month.split("-"))
    except Exception:
        raise HTTPException(400, "Use YYYY-MM format")
    from calendar import monthrange as _mr
    start = date(year, mon, 1)
    end   = date(year, mon, _mr(year, mon)[1])

    q = db.query(_WMEntry).filter(
        _WMEntry.entry_date >= start,
        _WMEntry.entry_date <= end,
    )
    if employee_id:
        q = q.filter(_WMEntry.employee_id == employee_id)
    if status:
        q = q.filter(_WMEntry.status == status)

    rows = q.order_by(_WMEntry.entry_date.asc()).all()
    return [_wm_hr_dict(r) for r in rows]


@router.put("/work-mode/{entry_id}/approve")
def hr_approve_work_mode(entry_id: int, data: _WMAction = _WMAction(), db: Session = Depends(get_db)):
    entry = db.query(_WMEntry).filter(_WMEntry.id == entry_id).first()
    if not entry:
        raise HTTPException(404, "Entry not found")
    entry.status = "Approved"
    if data.remarks:
        entry.hr_remarks = data.remarks
    entry.updated_at = _wm_dt.utcnow()
    db.commit()
    return {"ok": True, "status": "Approved"}


@router.put("/work-mode/{entry_id}/reject")
def hr_reject_work_mode(entry_id: int, data: _WMAction = _WMAction(), db: Session = Depends(get_db)):
    entry = db.query(_WMEntry).filter(_WMEntry.id == entry_id).first()
    if not entry:
        raise HTTPException(404, "Entry not found")
    entry.status = "Rejected"
    if data.remarks:
        entry.hr_remarks = data.remarks
    entry.updated_at = _wm_dt.utcnow()
    db.commit()
    return {"ok": True, "status": "Rejected"}


# ── Edit / Correction Requests (HR view) ──────────────────────

from backend.models.edit_request import EditRequest as _EditReq
from datetime import datetime as _er_dt


class _ResolveIn(BaseModel):
    hr_remarks: Optional[str] = None


@router.get("/edit-requests")
def hr_list_edit_requests(status: Optional[str] = None, db: Session = Depends(get_db)):
    q = db.query(_EditReq)
    if status and status != "All":
        q = q.filter(_EditReq.status == status)
    rows = q.order_by(_EditReq.created_at.desc()).all()
    return [
        {
            "id":            r.id,
            "employee_id":   r.employee_id,
            "employee_name": r.employee.full_name if r.employee else "—",
            "employee_code": r.employee.employee_id if r.employee else "—",
            "request_type":  r.request_type,
            "target_date":   str(r.target_date),
            "description":   r.description,
            "reason":        r.reason,
            "status":        r.status,
            "hr_remarks":    r.hr_remarks,
            "created_at":    str(r.created_at)[:10],
            "resolved_at":   str(r.resolved_at)[:10] if r.resolved_at else None,
        }
        for r in rows
    ]


@router.put("/edit-requests/{req_id}/approve")
def hr_approve_edit_request(req_id: int, request: Request, data: _ResolveIn = _ResolveIn(), db: Session = Depends(get_db)):
    req = db.query(_EditReq).filter(_EditReq.id == req_id).first()
    if not req:
        raise HTTPException(404, "Request not found")
    require_approval_rights(request, db, req.employee_id)
    req.status = "Approved"
    req.hr_remarks = data.hr_remarks
    req.resolved_at = _er_dt.utcnow()
    db.commit()
    return {"ok": True}


@router.put("/edit-requests/{req_id}/reject")
def hr_reject_edit_request(req_id: int, request: Request, data: _ResolveIn = _ResolveIn(), db: Session = Depends(get_db)):
    req = db.query(_EditReq).filter(_EditReq.id == req_id).first()
    if not req:
        raise HTTPException(404, "Request not found")
    require_approval_rights(request, db, req.employee_id)
    req.status = "Rejected"
    req.hr_remarks = data.hr_remarks
    req.resolved_at = _er_dt.utcnow()
    db.commit()
    return {"ok": True}


# ── Company Documents ──────────────────────────────────────────
import os as _os
from fastapi.responses import FileResponse as _FileResponse

COMPANY_DOCS_DIR = "/app/company_docs"

_ALLOWED_DOC_EXTS = {
    "pdf", "doc", "docx", "xls", "xlsx", "ppt", "pptx",
    "txt", "csv", "jpg", "jpeg", "png", "gif", "webp",
    "zip", "rar", "7z", "odt", "ods", "odp",
}

_MIME_MAP = {
    "pdf": "application/pdf",
    "doc": "application/msword",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xls": "application/vnd.ms-excel",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "ppt": "application/vnd.ms-powerpoint",
    "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "txt": "text/plain",
    "csv": "text/csv",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
    "gif": "image/gif",
    "webp": "image/webp",
    "zip": "application/zip",
    "rar": "application/x-rar-compressed",
    "7z": "application/x-7z-compressed",
    "odt": "application/vnd.oasis.opendocument.text",
    "ods": "application/vnd.oasis.opendocument.spreadsheet",
    "odp": "application/vnd.oasis.opendocument.presentation",
}


@router.get("/company-docs")
def list_company_docs(db: Session = Depends(get_db)):
    from sqlalchemy import text as _t
    # R2 is always the source of truth for existence
    files = storage.list_files("company-docs")
    if not files:
        return []
    # Backfill any R2 files not yet tracked in DB
    db_rows = {r[0] for r in db.execute(_t("SELECT name FROM company_documents")).fetchall()}
    for f in files:
        if f not in db_rows:
            db.execute(_t("""
                INSERT INTO company_documents (name, r2_key)
                VALUES (:name, :key) ON CONFLICT (name) DO NOTHING
            """), {"name": f, "key": f"company-docs/{f}"})
    db.commit()
    # Return ordered by upload time
    rows = db.execute(_t(
        "SELECT name, uploaded_by, uploaded_at FROM company_documents WHERE name = ANY(:names) ORDER BY uploaded_at DESC"
    ), {"names": files}).fetchall()
    return [{"name": r[0], "uploaded_by": r[1], "uploaded_at": str(r[2]) if r[2] else None,
             "url": f"/api/hrm/company-docs/{r[0]}"} for r in rows]


@router.get("/company-docs/{filename}")
def get_company_doc(filename: str):
    from fastapi.responses import Response as _Resp
    safe_name = _os.path.basename(filename)
    ext = safe_name.rsplit(".", 1)[-1].lower() if "." in safe_name else ""
    mime = _MIME_MAP.get(ext, "application/octet-stream")
    try:
        data = storage.download_file(f"company-docs/{safe_name}")
        return _Resp(content=data, media_type=mime,
                     headers={"Content-Disposition": f'attachment; filename="{safe_name}"'})
    except Exception:
        raise HTTPException(404, "File not found")


@router.post("/company-docs/upload", status_code=201)
async def upload_company_doc(request: Request, file: UploadFile = File(...), db: Session = Depends(get_db)):
    from sqlalchemy import text as _t
    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in _ALLOWED_DOC_EXTS:
        raise HTTPException(400, f"File type '.{ext}' is not allowed")
    safe_name = _os.path.basename(file.filename)
    if not safe_name:
        raise HTTPException(400, "Invalid filename")
    r2_key = f"company-docs/{safe_name}"
    storage.upload_file(await file.read(), "company-docs", safe_name)
    uploaded_by = getattr(request.state, "username", None)
    db.execute(_t("""
        INSERT INTO company_documents (name, r2_key, uploaded_by)
        VALUES (:name, :r2_key, :by)
        ON CONFLICT (name) DO UPDATE SET r2_key = EXCLUDED.r2_key, uploaded_by = EXCLUDED.uploaded_by, uploaded_at = NOW()
    """), {"name": safe_name, "r2_key": r2_key, "by": uploaded_by})
    db.commit()
    return {"name": safe_name, "url": f"/api/hrm/company-docs/{safe_name}"}


@router.delete("/company-docs/{filename}", status_code=200)
def delete_company_doc(filename: str, request: Request, db: Session = Depends(get_db)):
    from sqlalchemy import text as _t
    safe_name = _os.path.basename(filename)
    storage.delete_file(f"/files/company-docs/{safe_name}")
    deleted_by = getattr(request.state, "username", None)
    db.execute(_t("DELETE FROM company_documents WHERE name = :name"), {"name": safe_name})
    db.execute(_t("""
        INSERT INTO deletion_log (entity_type, entity_name, deleted_by, extra)
        VALUES ('company_document', :name, :by, :extra::jsonb)
    """), {"name": safe_name, "by": deleted_by, "extra": f'{{"r2_key":"company-docs/{safe_name}"}}'})
    db.commit()
    return {"deleted": safe_name}


# ── Letter generation ───────────────────────────────────────────────────────

from backend.utils.letter_generator import generate_letter, LETTER_FIELDS
from backend.models.employee import Employee as _Employee
from datetime import datetime as _dt
from fastapi.responses import Response as _Response
from pydantic import BaseModel as _BM
from typing import Any as _Any

GENERATED_DOCS_DIR = "/app/generated_docs"


class LetterGenerateRequest(_BM):
    letter_type: str
    employee_id: int
    fields: dict[str, _Any]


@router.get("/letter-fields")
def get_letter_fields():
    """Return field definitions for each letter type."""
    return LETTER_FIELDS


@router.post("/letters/generate")
def generate_employee_letter(
    body: LetterGenerateRequest,
    db: Session = Depends(get_db),
):
    emp = db.query(_Employee).filter(_Employee.id == body.employee_id).first()
    if not emp:
        raise HTTPException(404, "Employee not found")

    # Merge employee data into fields
    fields = dict(body.fields)
    fields["employee_name"] = emp.full_name
    fields["designation"]   = fields.get("designation") or (emp.designation_rel.name if emp.designation_rel else "")
    fields["employee_code"] = emp.employee_id or ""
    if not fields.get("department") and emp.department_rel:
        fields["department"] = emp.department_rel.name

    # Load active letterhead template config
    from backend.models.hrm import LetterheadTemplate as _LT
    tpl_row = db.query(_LT).filter(_LT.id == 1).first()
    tpl_cfg = {
        c.key: getattr(tpl_row, c.key)
        for c in _LT.__table__.columns
        if c.key not in ("id", "updated_at")
    } if tpl_row else {}

    try:
        pdf_bytes = generate_letter(body.letter_type, fields, template=tpl_cfg)
    except ValueError as e:
        raise HTTPException(400, str(e))

    # Save generated PDF so it's available for download + employee record
    safe_type = body.letter_type.replace(" ", "_").replace("/", "_")
    timestamp = _dt.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"{emp.employee_id}_{safe_type}_{timestamp}.pdf"
    storage.upload_file(pdf_bytes, "generated-letters", filename)

    # Create a fulfilled DocumentRequest so employee sees it in My Documents
    from backend.models.document_request import DocumentRequest as _DR
    doc_req = _DR(
        employee_id=body.employee_id,
        doc_type=body.letter_type,
        status="Fulfilled",
        requested_at=_dt.utcnow(),
        fulfilled_at=_dt.utcnow(),
        file_url=f"/api/hrm/letters/download/{filename}",
        file_name=filename,
    )
    db.add(doc_req)
    db.commit()

    return _Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/letters/download/{filename}")
def download_generated_letter(filename: str):
    from fastapi.responses import Response as _Resp
    safe_name = _os.path.basename(filename)
    try:
        data = storage.download_file(f"generated-letters/{safe_name}")
        return _Resp(
            content=data,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
        )
    except Exception:
        raise HTTPException(404, "File not found")


# ── Custom Document Templates ────────────────────────────────────────────────
from backend.models.doc_template import DocumentTemplate as _DocTpl


class _DocTplCreate(_BM):
    name: str
    category: str = "HR Letter"
    content: str
    variables: list[dict] = []


class _DocTplUpdate(_BM):
    name: str | None = None
    category: str | None = None
    content: str | None = None
    variables: list[dict] | None = None


class _DocTplGenerateReq(_BM):
    employee_id: int | None = None
    fields: dict[str, _Any] = {}


@router.get("/doc-templates")
def list_doc_templates(db: Session = Depends(get_db)):
    return db.query(_DocTpl).order_by(_DocTpl.created_at.desc()).all()


@router.post("/doc-templates")
def create_doc_template(body: _DocTplCreate, db: Session = Depends(get_db)):
    tpl = _DocTpl(**body.dict())
    db.add(tpl)
    db.commit()
    db.refresh(tpl)
    return tpl


@router.put("/doc-templates/{tpl_id}")
def update_doc_template(tpl_id: int, body: _DocTplUpdate, db: Session = Depends(get_db)):
    tpl = db.query(_DocTpl).filter(_DocTpl.id == tpl_id).first()
    if not tpl:
        raise HTTPException(404, "Template not found")
    for k, v in body.dict(exclude_unset=True).items():
        setattr(tpl, k, v)
    tpl.updated_at = _dt.utcnow()
    db.commit()
    db.refresh(tpl)
    return tpl


@router.delete("/doc-templates/{tpl_id}")
def delete_doc_template(tpl_id: int, db: Session = Depends(get_db)):
    tpl = db.query(_DocTpl).filter(_DocTpl.id == tpl_id).first()
    if not tpl:
        raise HTTPException(404, "Template not found")
    db.delete(tpl)
    db.commit()
    return {"ok": True}


@router.post("/doc-templates/{tpl_id}/generate")
def generate_from_doc_template(
    tpl_id: int,
    body: _DocTplGenerateReq,
    db: Session = Depends(get_db),
):
    tpl = db.query(_DocTpl).filter(_DocTpl.id == tpl_id).first()
    if not tpl:
        raise HTTPException(404, "Template not found")

    fields = dict(body.fields)

    emp = None
    if body.employee_id:
        emp = db.query(_Employee).filter(_Employee.id == body.employee_id).first()
        if emp:
            fields.setdefault("employee_name", emp.full_name)
            fields.setdefault("candidate_name", emp.full_name)
            fields.setdefault("employee_id_code", emp.employee_id or "")
            fields.setdefault("designation", emp.designation_rel.name if emp.designation_rel else "")
            fields.setdefault("department", emp.department_rel.name if emp.department_rel else "")
            fields.setdefault("work_email", emp.email or "")

    from backend.models.hrm import LetterheadTemplate as _LT
    tpl_row = db.query(_LT).filter(_LT.id == 1).first()
    tpl_cfg = {
        col.key: getattr(tpl_row, col.key)
        for col in _LT.__table__.columns
        if col.key not in ("id", "updated_at")
    } if tpl_row else {}

    from backend.utils.letter_generator import generate_custom_letter
    pdf_bytes = generate_custom_letter(tpl.content, fields, template=tpl_cfg)

    safe_name = tpl.name.replace(" ", "_").replace("/", "_")
    timestamp  = _dt.utcnow().strftime("%Y%m%d_%H%M%S")
    emp_prefix = (emp.employee_id or str(body.employee_id)) if emp else "doc"
    filename   = f"{emp_prefix}_{safe_name}_{timestamp}.pdf"
    storage.upload_file(pdf_bytes, "generated-letters", filename)

    if emp:
        from backend.models.document_request import DocumentRequest as _DR
        db.add(_DR(
            employee_id=body.employee_id,
            doc_type=tpl.name,
            status="Fulfilled",
            requested_at=_dt.utcnow(),
            fulfilled_at=_dt.utcnow(),
            file_url=f"/api/hrm/letters/download/{filename}",
            file_name=filename,
        ))
        db.commit()

    return _Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "X-Filename": filename,
        },
    )


class _ExtractFromDocReq(_BM):
    doc_name: str  # filename from /api/hrm/company-docs list


def _normalise_placeholders(text: str):
    """Convert all placeholder formats to {{snake_case}}.

    Handles:
      {{Employee Name}}  →  {{employee_name}}   (spaced / Title Case)
      {{EmployeeName}}   →  {{employeename}}     (no separator)
      [Employee Name]    →  {{employee_name}}    (square brackets)
      {{joining_date}}   →  unchanged            (already snake_case)

    Returns (converted_text, variables_list).
    """
    import re as _re

    label_map: dict[str, str] = {}  # snake_key → original display label

    def _to_snake(label: str) -> str:
        key = _re.sub(r'\s+', '_', label.strip().lower())
        key = _re.sub(r'[^a-z0-9_]', '', key)
        return key

    def curly_replacer(m):
        inner = m.group(1)
        # Already a valid snake_case identifier — keep as-is
        if _re.match(r'^[a-z][a-z0-9_]*$', inner):
            return m.group(0)
        label = inner.strip()
        key = _to_snake(label)
        if not key:
            return m.group(0)
        label_map[key] = label   # remember original label for display
        return f"{{{{{key}}}}}"

    def bracket_replacer(m):
        label = m.group(1).strip()
        key = _to_snake(label)
        if not key:
            return m.group(0)
        label_map[key] = label
        return f"{{{{{key}}}}}"

    # Step 1: normalize {{Any Format}} → {{snake_case}}
    converted = _re.sub(r'\{\{([^}]+)\}\}', curly_replacer, text)

    # Step 2: convert [Square Bracket Placeholder] → {{snake_case}}
    converted = _re.sub(r'\[([^\[\]]{1,60})\]', bracket_replacer, converted)

    # Step 3: collect all {{keys}} in document order (deduplicated)
    all_keys = list(dict.fromkeys(_re.findall(r'\{\{(\w+)\}\}', converted)))
    variables = []
    for k in all_keys:
        original_label = label_map.get(k, k.replace("_", " ").title())
        variables.append({
            "key":   k,
            "label": original_label,
            "type":  "date" if any(w in k for w in ("date", "dob", "doj")) else "text",
        })

    return converted, variables


@router.post("/doc-templates/extract-from-doc")
def extract_from_existing_doc(body: _ExtractFromDocReq):
    """Extract text + detect variables from an already-uploaded company document."""
    safe = _os.path.basename(body.doc_name)

    try:
        file_data = storage.download_file(f"company-docs/{safe}")
    except Exception:
        raise HTTPException(404, "Document not found in company docs")

    ext = safe.rsplit(".", 1)[-1].lower() if "." in safe else ""
    text = ""

    if ext == "docx":
        try:
            from docx import Document as _DocxDoc
            import io as _io
            doc = _DocxDoc(_io.BytesIO(file_data))
            text = "\n".join(p.text for p in doc.paragraphs).strip()
        except Exception as e:
            raise HTTPException(400, f"Could not read DOCX: {e}")
    elif ext == "pdf":
        try:
            from pypdf import PdfReader as _PdfReader
            import io as _io
            reader = _PdfReader(_io.BytesIO(file_data))
            text = "\n\n".join(p.extract_text() or "" for p in reader.pages).strip()
        except Exception as e:
            raise HTTPException(400, f"Could not read PDF: {e}")
    else:
        raise HTTPException(400, "Only PDF and DOCX documents are supported")

    if not text:
        raise HTTPException(422, "No selectable text found — document may be a scanned image")

    converted_text, variables = _normalise_placeholders(text)
    return {"text": converted_text, "variables": variables}


@router.post("/doc-templates/extract-text")
async def extract_text_from_file(file: UploadFile = File(...)):
    """Extract plain text from an uploaded PDF or DOCX so it can be used as template content."""
    fname = (file.filename or "").lower()
    data  = await file.read()

    if fname.endswith(".docx"):
        try:
            from docx import Document as _DocxDoc
            import io as _io
            doc = _DocxDoc(_io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs).strip()
        except Exception as e:
            raise HTTPException(400, f"Could not read DOCX: {e}")

    elif fname.endswith(".pdf"):
        try:
            from pypdf import PdfReader as _PdfReader
            import io as _io
            reader = _PdfReader(_io.BytesIO(data))
            text   = "\n\n".join(p.extract_text() or "" for p in reader.pages).strip()
        except Exception as e:
            raise HTTPException(400, f"Could not read PDF: {e}")

    else:
        raise HTTPException(400, "Only PDF and DOCX files are supported")

    if not text:
        raise HTTPException(422, "No text could be extracted from this file. Make sure the document contains selectable text (not a scanned image).")

    converted_text, variables = _normalise_placeholders(text)
    return {"text": converted_text, "variables": variables}


# ── Letterhead Template ──────────────────────────────────────────────────────
from backend.models.hrm import LetterheadTemplate as _LetterheadTemplate

_LOGO_UPLOADS_DIR = "/app/letterhead_logos"

_LH_FIELDS = [
    "company_name", "tagline", "logo_filename",
    "logo_x_mm", "logo_y_mm", "logo_w_mm", "logo_h_mm",
    "footer_image_filename",
    "footer_x_mm", "footer_y_mm", "footer_w_mm", "footer_h_mm",
    "signature_filename", "sig_x_mm", "sig_w_mm", "sig_h_mm",
    "addr1", "addr2", "addr3", "addr4",
    "phone1", "phone2", "email", "website",
    "header_color", "accent_color",
    "hr_signatory", "hr_role",
    "content_top_mm", "body_font", "body_font_size", "body_bold", "body_italic",
    "watermark_filename", "watermark_opacity",
    "watermark_x_mm", "watermark_y_mm", "watermark_w_mm", "watermark_h_mm",
]
_FOOTER_UPLOADS_DIR = "/app/letterhead_logos"


def _get_or_create_template(db):
    row = db.query(_LetterheadTemplate).filter(_LetterheadTemplate.id == 1).first()
    if not row:
        row = _LetterheadTemplate(id=1)
        db.add(row)
        db.commit()
        db.refresh(row)
    return row


@router.get("/letterhead-template")
def get_letterhead_template(db: Session = Depends(get_db)):
    row = _get_or_create_template(db)
    return {f: getattr(row, f) for f in _LH_FIELDS}


class _LHUpdate(_BM):
    company_name: str | None = None
    tagline: str | None = None
    logo_x_mm: float | None = None
    logo_y_mm: float | None = None
    logo_w_mm: float | None = None
    logo_h_mm: float | None = None
    footer_x_mm: float | None = None
    footer_y_mm: float | None = None
    footer_w_mm: float | None = None
    footer_h_mm: float | None = None
    sig_x_mm: float | None = None
    sig_w_mm: float | None = None
    sig_h_mm: float | None = None
    addr1: str | None = None
    addr2: str | None = None
    addr3: str | None = None
    addr4: str | None = None
    phone1: str | None = None
    phone2: str | None = None
    email: str | None = None
    website: str | None = None
    header_color: str | None = None
    accent_color: str | None = None
    hr_signatory: str | None = None
    hr_role: str | None = None
    content_top_mm: float | None = None
    body_font: str | None = None
    body_font_size: float | None = None
    body_bold: bool | None = None
    body_italic: bool | None = None
    watermark_opacity: float | None = None
    watermark_x_mm: float | None = None
    watermark_y_mm: float | None = None
    watermark_w_mm: float | None = None
    watermark_h_mm: float | None = None


@router.put("/letterhead-template")
def update_letterhead_template(body: _LHUpdate, db: Session = Depends(get_db)):
    row = _get_or_create_template(db)
    for field, val in body.model_dump(exclude_none=True).items():
        setattr(row, field, val)
    db.commit()
    db.refresh(row)
    return {f: getattr(row, f) for f in _LH_FIELDS}


@router.post("/letterhead-template/logo")
async def upload_letterhead_logo(file: UploadFile = File(...), db: Session = Depends(get_db)):
    allowed = ('.png', '.jpg', '.jpeg', '.webp')
    ext = _os.path.splitext(file.filename or '')[1].lower()
    if ext not in allowed:
        raise HTTPException(400, "Only PNG/JPG/WEBP images are allowed")
    filename = f"logo{ext}"
    contents = await file.read()
    storage.upload_file(contents, "letterhead", filename)
    row = _get_or_create_template(db)
    row.logo_filename = filename
    db.commit()
    return {"logo_filename": filename, "url": f"/api/hrm/letterhead-template/logo/{filename}"}


@router.get("/letterhead-template/logo/{filename}")
def get_letterhead_logo(filename: str):
    from fastapi.responses import Response as _Resp
    safe = _os.path.basename(filename)
    try:
        data = storage.download_file(f"letterhead/{safe}")
    except Exception:
        raise HTTPException(404, "Logo not found")
    ext = _os.path.splitext(safe)[1].lower()
    mt = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp"}.get(ext.lstrip('.'), "application/octet-stream")
    return _Resp(content=data, media_type=mt)


@router.delete("/letterhead-template/logo", status_code=200)
def delete_letterhead_logo(db: Session = Depends(get_db)):
    row = _get_or_create_template(db)
    if row.logo_filename:
        storage.delete_file(f"/files/letterhead/{row.logo_filename}")
        row.logo_filename = None
        db.commit()
    return {"ok": True}


@router.post("/letterhead-template/footer-image")
async def upload_footer_image(file: UploadFile = File(...), db: Session = Depends(get_db)):
    allowed = ('.png', '.jpg', '.jpeg', '.webp')
    ext = _os.path.splitext(file.filename or '')[1].lower()
    if ext not in allowed:
        raise HTTPException(400, "Only PNG/JPG/WEBP images are allowed")
    filename = f"footer{ext}"
    contents = await file.read()
    storage.upload_file(contents, "letterhead", filename)
    row = _get_or_create_template(db)
    row.footer_image_filename = filename
    db.commit()
    return {"footer_image_filename": filename, "url": f"/api/hrm/letterhead-template/footer-image/{filename}"}


@router.get("/letterhead-template/footer-image/{filename}")
def get_footer_image(filename: str):
    from fastapi.responses import Response as _Resp
    safe = _os.path.basename(filename)
    try:
        data = storage.download_file(f"letterhead/{safe}")
    except Exception:
        raise HTTPException(404, "Footer image not found")
    ext = _os.path.splitext(safe)[1].lower()
    mt = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp"}.get(ext.lstrip('.'), "application/octet-stream")
    return _Resp(content=data, media_type=mt)


@router.delete("/letterhead-template/footer-image", status_code=200)
def delete_footer_image(db: Session = Depends(get_db)):
    row = _get_or_create_template(db)
    if row.footer_image_filename:
        storage.delete_file(f"/files/letterhead/{row.footer_image_filename}")
        row.footer_image_filename = None
        db.commit()
    return {"ok": True}


@router.post("/letterhead-template/signature")
async def upload_signature(file: UploadFile = File(...), db: Session = Depends(get_db)):
    allowed = ('.png', '.jpg', '.jpeg', '.webp')
    ext = _os.path.splitext(file.filename or '')[1].lower()
    if ext not in allowed:
        raise HTTPException(400, "Only PNG/JPG/WEBP images are allowed")
    filename = f"signature{ext}"
    storage.upload_file(await file.read(), "letterhead", filename)
    row = _get_or_create_template(db)
    row.signature_filename = filename
    db.commit()
    return {"signature_filename": filename, "url": f"/api/hrm/letterhead-template/signature/{filename}"}


@router.get("/letterhead-template/signature/{filename}")
def get_signature(filename: str):
    from fastapi.responses import Response as _Resp
    safe = _os.path.basename(filename)
    try:
        data = storage.download_file(f"letterhead/{safe}")
    except Exception:
        raise HTTPException(404, "Signature not found")
    ext = _os.path.splitext(safe)[1].lower()
    mt = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp"}.get(ext.lstrip('.'), "application/octet-stream")
    return _Resp(content=data, media_type=mt)


@router.delete("/letterhead-template/signature", status_code=200)
def delete_signature(db: Session = Depends(get_db)):
    row = _get_or_create_template(db)
    if row.signature_filename:
        storage.delete_file(f"/files/letterhead/{row.signature_filename}")
        row.signature_filename = None
        db.commit()
    return {"ok": True}


@router.post("/letterhead-template/watermark")
async def upload_watermark(file: UploadFile = File(...), db: Session = Depends(get_db)):
    allowed = ('.png', '.jpg', '.jpeg', '.webp')
    ext = _os.path.splitext(file.filename or '')[1].lower()
    if ext not in allowed:
        raise HTTPException(400, "Only PNG/JPG/WEBP images are allowed")
    filename = f"watermark{ext}"
    storage.upload_file(await file.read(), "letterhead", filename)
    row = _get_or_create_template(db)
    row.watermark_filename = filename
    db.commit()
    return {"watermark_filename": filename, "url": f"/api/hrm/letterhead-template/watermark/{filename}"}


@router.get("/letterhead-template/watermark/{filename}")
def get_watermark(filename: str):
    from fastapi.responses import Response as _Resp
    safe = _os.path.basename(filename)
    try:
        data = storage.download_file(f"letterhead/{safe}")
    except Exception:
        raise HTTPException(404, "Watermark not found")
    ext = _os.path.splitext(safe)[1].lower()
    mt = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp"}.get(ext.lstrip('.'), "application/octet-stream")
    return _Resp(content=data, media_type=mt)


@router.delete("/letterhead-template/watermark", status_code=200)
def delete_watermark(db: Session = Depends(get_db)):
    row = _get_or_create_template(db)
    if row.watermark_filename:
        storage.delete_file(f"/files/letterhead/{row.watermark_filename}")
        row.watermark_filename = None
        db.commit()
    return {"ok": True}


@router.post("/letterhead-template/preview")
def preview_letterhead_template(db: Session = Depends(get_db)):
    """Generate a sample Appointment Letter using the saved template settings."""
    row = _get_or_create_template(db)
    tpl_cfg = {f: getattr(row, f) for f in _LH_FIELDS}
    sample_fields = {
        "employee_name":   "Ravi Kumar",
        "designation":     "Software Engineer",
        "department":      "Engineering",
        "letter_date":     "2025-01-15",
        "joining_date":    "2024-10-01",
        "confirmation_date": "2025-01-01",
    }
    pdf_bytes = generate_letter("Appointment Letter", sample_fields, template=tpl_cfg)
    return _Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": 'inline; filename="template-preview.pdf"'},
    )
